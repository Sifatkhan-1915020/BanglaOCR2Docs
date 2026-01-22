import streamlit as st
import pytesseract
from pdf2image import convert_from_bytes
from docx import Document
import io
import unicodedata
import platform

# --- 1. PAGE CONFIGURATION (Must be first) ---
st.set_page_config(
    page_title="Bangla OCR Pro",
    page_icon="🇧🇩",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 2. MODERN CSS STYLING ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Hind+Siliguri:wght@300;400;600&display=swap');

    /* Apply Font to Whole App */
    html, body, [class*="css"] {
        font-family: 'Hind Siliguri', sans-serif;
    }

    /* Main Title Styling */
    .main-title {
        font-size: 3rem;
        font-weight: 700;
        background: linear-gradient(to right, #006a4e, #f42a41); /* BD Flag Colors Gradient */
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        margin-bottom: 0px;
    }

    /* Subtitle */
    .subtitle {
        font-size: 1.2rem;
        color: #555;
        margin-bottom: 30px;
    }

    /* Modern Card Container */
    .card {
        background-color: #f9f9f9;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 20px;
    }

    /* Gradient Button */
    div.stButton > button:first-child {
        background: linear-gradient(45deg, #006a4e, #2e8b57);
        color: white;
        border: none;
        padding: 0.6rem 2rem;
        border-radius: 50px;
        font-size: 1rem;
        font-weight: bold;
        transition: all 0.3s ease;
        width: 100%;
    }
    div.stButton > button:first-child:hover {
        transform: scale(1.02);
        box-shadow: 0 6px 12px rgba(0,106,78,0.3);
    }

    /* Text Area Styling */
    .stTextArea textarea {
        background-color: #ffffff;
        border: 1px solid #ddd;
        border-radius: 10px;
        font-family: 'Hind Siliguri', sans-serif !important;
        font-size: 1.1rem;
    }
</style>
""", unsafe_allow_html=True)

# --- 3. FUNCTIONS (Your Logic) ---

def extract_text_from_pdf(pdf_bytes):
    """
    Extracts text using Tesseract + Poppler + Unicode Fix
    """
    # ⚠️ CHECK YOUR POPPLER PATH
    

    try:
        if platform.system() == "Windows":
         poppler_path = r"C:\Program Files\poppler\Library\bin"
         images = convert_from_bytes(pdf_bytes, poppler_path=poppler_path)
        else:
        # On Linux (Streamlit Cloud), Poppler is in the global path, so we don't need to specify it.
         images = convert_from_bytes(pdf_bytes)
    except Exception as e:
        st.error(f"❌ Poppler Error: {e}")
        return None

    extracted_text = ""
    total_pages = len(images)
    
    # Create a placeholder for progress
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i, img in enumerate(images):
        # Update UI
        status_text.markdown(f"**Processing Page {i+1} of {total_pages}...**")
        progress_bar.progress(int(((i + 1) / total_pages) * 100))
        
        # OCR Configuration (Best for Bangla)
        custom_config = r'--oem 1 --psm 6'
        text = pytesseract.image_to_string(img, lang='ben+eng', config=custom_config)
        
        # Unicode Normalization (The Fix)
        fixed_text = unicodedata.normalize('NFC', text)
        extracted_text += fixed_text + "\n\n"
    
    # Cleanup UI
    progress_bar.empty()
    status_text.empty()
    return extracted_text

def create_word_docx(text):
    doc = Document()
    doc.add_heading('Converted Bangla Document', 0)
    
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip(): 
            p = doc.add_paragraph(para)
            p.style.font.name = 'Arial'
            
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. MAIN APP LAYOUT ---

# Sidebar for Info
with st.sidebar:
    st.image("https://upload.wikimedia.org/wikipedia/commons/thumb/f/fa/Flag_of_the_People%27s_Republic_of_Bangladesh.svg/2560px-Flag_of_the_People%27s_Republic_of_Bangladesh.svg.png", width=50)
    st.header("Bangla OCR Pro")
    st.markdown("Convert scanned PDFs to editable Word documents using advanced **Tesseract LSTM** engine.")
    st.markdown("---")
    st.caption("✅ **Features:**")
    st.markdown("""
    - 🔒 **100% Offline** (Privacy First)
    - 🛠️ **Unicode Fixer** (No broken letters)
    - 📄 **Word Export** (.docx)
    """)
    st.markdown("---")
    st.caption("Powered by Python & Tesseract")

# Main Content Area
col1, col2 = st.columns([1, 2])

with col1:
    st.markdown('<p class="main-title">Start Here</p>', unsafe_allow_html=True)
    st.markdown('<div class="card">', unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader("📂 Upload your Bangla PDF", type="pdf")
    
    if uploaded_file is not None:
        st.success(f"File Uploaded: {uploaded_file.name}")
        
        if st.button("🚀 Convert Now"):
            with st.spinner("🧠 AI is reading your document..."):
                pdf_bytes = uploaded_file.read()
                
                # Store result in session state to prevent reload loss
                st.session_state['final_text'] = extract_text_from_pdf(pdf_bytes)
                st.session_state['file_name'] = uploaded_file.name
    
    st.markdown('</div>', unsafe_allow_html=True)

with col2:
    st.markdown('<p class="main-title">Results</p>', unsafe_allow_html=True)
    
    if 'final_text' in st.session_state:
        text_data = st.session_state['final_text']
        
        if text_data:
            # Result Card
            st.markdown('<div class="card">', unsafe_allow_html=True)
            
            tab1, tab2 = st.tabs(["📝 Text Preview", "💾 Download"])
            
            with tab1:
                st.markdown("### extracted content:")
                st.text_area("Editable Preview", text_data, height=500, label_visibility="collapsed")
            
            with tab2:
                st.markdown("### Ready to Export?")
                st.write("Your document has been corrected and formatted.")
                
                docx_file = create_word_docx(text_data)
                
                # Create a filename
                out_name = st.session_state['file_name'].replace('.pdf', '_converted.docx')
                
                st.download_button(
                    label="📥 Download Word Document (.docx)",
                    data=docx_file,
                    file_name=out_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.markdown('</div>', unsafe_allow_html=True)
    else:
        # Empty State Placeholder
        st.markdown("""
        <div style="text-align: center; color: #aaa; margin-top: 50px;">
            <h3>👈 Upload a file to see magic happen</h3>
            <p>Your results will appear here automatically.</p>
        </div>
        """, unsafe_allow_html=True)